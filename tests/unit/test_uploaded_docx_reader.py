"""上传 DOCX 读取工具单元测试。"""

import json
from io import BytesIO
from types import SimpleNamespace

from docx import Document

from app.ai.tools import file_tools


def _build_docx_bytes() -> bytes:
    document = Document()
    document.add_heading("DOCX回归测试", level=1)
    document.add_paragraph("第一段正文")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "列1"
    table.cell(0, 1).text = "列2"
    table.cell(1, 0).text = "值A"
    table.cell(1, 1).text = "值B"
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


class _Resp:
    def __init__(self, content: bytes):
        self._content = content

    def read(self) -> bytes:
        return self._content

    def close(self) -> None:
        return None

    def release_conn(self) -> None:
        return None


def _install_asset_service(monkeypatch, *, content: bytes, content_type: str, size: int | None = None) -> None:
    stat = SimpleNamespace(content_type=content_type, size=size if size is not None else len(content))
    asset_service = SimpleNamespace(
        client=SimpleNamespace(
            get_object=lambda **kwargs: _Resp(content),
            stat_object=lambda **kwargs: stat,
        )
    )
    monkeypatch.setattr(file_tools, "get_asset_service", lambda: asset_service)


def test_read_uploaded_file_should_extract_docx_text(monkeypatch) -> None:
    payload = _build_docx_bytes()
    _install_asset_service(
        monkeypatch,
        content=payload,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    result = file_tools.read_uploaded_file.func(
        file_path="/api/v1/assets/1/thread/sample.docx",
        config={},
    )
    data = json.loads(result)

    assert data["status"] == "success"
    assert data["file_type"] == "docx"
    assert "DOCX回归测试" in data["content"]
    assert "第一段正文" in data["content"]
    assert "值A" in data["content"]


def test_read_uploaded_file_should_report_invalid_docx(monkeypatch) -> None:
    _install_asset_service(
        monkeypatch,
        content=b"not-a-docx",
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    result = file_tools.read_uploaded_file.func(
        file_path="/api/v1/assets/1/thread/broken.docx",
        config={},
    )
    data = json.loads(result)

    assert data["status"] == "error"
    assert "DOCX" in data["message"]
    assert data["file_name"] == "broken.docx"


def test_read_uploaded_file_should_ask_to_convert_legacy_doc(monkeypatch) -> None:
    _install_asset_service(
        monkeypatch,
        content=b"legacy-doc",
        content_type="application/msword",
    )

    result = file_tools.read_uploaded_file.func(
        file_path="/api/v1/assets/1/thread/legacy.doc",
        config={},
    )
    data = json.loads(result)

    assert data["status"] == "unsupported"
    assert data["message"] == "暂不支持 .doc，请先转换为 .docx"
    assert data["object_key"].endswith("legacy.doc")
